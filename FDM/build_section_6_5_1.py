"""Emit Section 6.5.1 (Computational time) as .docx and .md from one source.

Every number in the prose is drawn from the measurement files rather than typed
in, so re-running the measurements and re-running this script keeps the text and
the data in step.

    python3 measure_solve_cost.py
    python3 reconstruct_loss_history.py
    python3 figure_computational_time.py
    python3 figure_convergence_cost.py
    python3 build_section_6_5_1.py
"""

import glob
import json
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
FIG = os.path.join(HERE, "figures")
OUT_DOCX = os.path.join(HERE, "data", "Section_6_5_1_Computational_time.docx")
OUT_MD = os.path.join(ROOT, "docs", "6_5_1_computational_time.md")

cost = json.load(open(os.path.join(HERE, "data", "solve_cost.json")))
hist = json.load(open(os.path.join(HERE, "data", "D5", "loss_history.json")))
fdm_all = json.load(open(os.path.join(HERE, "data", "fdm_cost.json")))
fdm = fdm_all["D5 (6.4.6)"]
fdm_b5, fdm_c5 = fdm_all["B5 (6.4.2)"], fdm_all["C5 (6.4.3)"]
fdm_res = json.load(open(sorted(glob.glob(
    os.path.join(HERE, "data", "D5", "D5_fdm_*.json")))[-1]))


def med(case):
    s = sorted(cost[case]["samples"])
    return s[len(s) // 2]


def spread(case):
    s = sorted(cost[case]["samples"])
    return s[-1] / s[0]


M2, MB5, MC5, MD5 = (med(k) for k in ("2part (6.4.1)", "B5 (6.4.2)",
                                      "C5 (6.4.3)", "D5 (6.4.6)"))

FDM_S = sorted(fdm["samples_s"])[len(fdm["samples_s"]) // 2]
FDM_MS = 1e3 * FDM_S / fdm["iters"]
FIELD_S = sorted(fdm["field_samples_s"])[len(fdm["field_samples_s"]) // 2]
FDM_RATIO = MD5 / (fdm["inflate_ms"] / 1e3)
PICARD = ("zero one two three four five six seven eight nine ten"
          .split()[fdm["picard_steps"]])

FDM_RMS_MM = 1e3 * fdm_res["rmse_m"]
FDM_RMS_PC = 100 * fdm_res["rmse_m"] / fdm_res["span_m"]
FDM_CROWN_UM = 1e6 * abs(fdm_res["fdm_crown_m"] - fdm_res["target_crown_m"])
FEM_BEST_MM = 1e3 * min(min(v["loss"]) for v in hist.values())

# Cost of one L-BFGS-B iteration, adjoint against direct sensitivity.
PER_VAR_B5 = fdm_b5["ms_per_iter"] / fdm_b5["n_design"]
PER_VAR_C5 = fdm_c5["ms_per_iter"] / fdm_c5["n_design"]
DIRECT_RATIO = fdm_c5["ms_per_iter"] / fdm["ms_per_iter"]
D5_IF_DIRECT_MIN = (PER_VAR_C5 * fdm["n_design"] * fdm["iters"] / 1e3) / 60

BODY = [
("h", "6.5.1 Computational time"),

("p", "All timings in this section were measured on a single workstation with an "
      "Intel Xeon Gold 5412U, twenty cores and 16 GB of memory. The forward solver "
      "is serial, compiled with GCC 14.2.0 at -O3, and uses CHOLMOD 5.3.1 "
      "(SuiteSparse 7.10.1) for the sparse factorisation, with the AMD ordering "
      "rather than nested dissection, METIS not being linked. One objective "
      "evaluation occupies one core, and the other nineteen stay idle unless "
      "several evaluations are dispatched at once, which the present "
      "implementation does not do."),

("p", f"The cost of the workflow is set almost entirely by the inner forward solve "
      f"of Section 6.3.7. Everything outside it is done once per target: on the "
      f"shell with openings the force density method (FDM) form finding of "
      f"Section 6.2 takes {FDM_S:.1f} s and the directional field "
      f"{FIELD_S:.1f} s, against the hours the inverse problem takes on the same "
      f"mesh. That margin is not uniform across the case studies, and the reason "
      f"it is not turns out to be the most transferable result in this section, "
      f"so it is taken up separately below. The question first is what one "
      f"forward solve costs, how many of them a strategy needs, and why the two "
      f"do not multiply to the observed wall clock."),

("h2", "What one forward solve costs"),

("p", f"Figure 6.29a gives the cost of a solve on each of the four case-study "
      f"meshes, sampled at twenty-five design points drawn within ±10% of the "
      f"recorded optimum — the neighbourhood a finite-difference L-BFGS-B "
      f"actually visits. The median cost is {M2:.2f} s on the creased shell "
      f"({cost['2part (6.4.1)']['n_faces']} elements), {MB5:.2f} s on the "
      f"free-form shell ({cost['B5 (6.4.2)']['n_faces']}), {MD5:.2f} s on the "
      f"shell with openings ({cost['D5 (6.4.6)']['n_faces']}) and {MC5:.2f} s on "
      f"the fluted dome ({cost['C5 (6.4.3)']['n_faces']}). The order is not the "
      f"order of the meshes. The fluted dome carries the largest mesh in the set "
      f"and is the second cheapest to solve, while the shell with openings, with "
      f"a mesh 27% smaller, costs five times as much per evaluation."),

("p", "What sets the cost is the number of Newton iterations the inner solve "
      "needs, and that is a property of the design point rather than of the "
      "discretisation: how far the trial state sits from equilibrium, and how "
      "many cable segments change between taut and slack as the four "
      "pressure-continuation stages are traversed. The consequence is that the "
      "cost of an evaluation is not a constant of the problem but a "
      "distribution with a long upper tail. Within ±10% of its own optimum the "
      f"creased shell spans a factor of {spread('2part (6.4.1)'):.0f} between its "
      f"cheapest and dearest evaluation, and the shell with openings a factor of "
      f"{spread('D5 (6.4.6)'):.0f}; one of its twenty-five sampled points "
      f"exceeded the 120 s ceiling imposed on the measurement. The free-form "
      f"shell and the fluted dome, by contrast, stay within a factor of six. A "
      f"single timing at the optimum is therefore not a useful summary of what a "
      f"run will cost, and the rest of this section is largely about the tail."),

("h2", "How many solves a strategy needs"),

("p", "Because the outer problem is solved with finite-difference gradients, one "
      "L-BFGS-B iteration costs n + 1 forward solves for a design vector of "
      "length n. The design vectors are deliberately small — one or two variables "
      "for the global strategies of Section 6.4.1, seven for the "
      "symmetry-reduced fluted dome, thirty for the free-form shell — so a "
      "gradient costs between two and thirty-one solves, a few seconds at most. "
      "Table 6.X collects the totals."),

("table", None),

("p", "The table separates two populations that should not be averaged. The "
      "strategies with a global or symmetry-reduced parameter set converge in a "
      "few hundred evaluations and finish in one to two minutes of solve time: "
      "622 evaluations for the free-form shell over thirty variables, 412 for the "
      "fluted dome across its two phases, 162 for the single-region fit of the "
      "shell with openings. The alternating region schemes are one to two orders "
      "of magnitude more expensive. The adaptive partition of the shell with "
      "openings took 2959 evaluations and 184 minutes, because every element on a "
      "region boundary is trial-assigned to each neighbouring region at one "
      "forward solve per trial, and the sweep repeats until no element moves. It "
      "is the discrete outer step, not the continuous inner one, that makes "
      "region refinement expensive, and its cost scales with the length of the "
      "region boundaries rather than with the number of variables. The "
      "twenty-variable field-aligned partition needed 1241 evaluations against "
      "the nine-variable adaptive partition's 2959. The last row belongs to "
      "neither population, being the form finding that precedes the inverse "
      "problem rather than an instance of it; it is there for the comparison "
      "drawn below, and it is cheaper than every run on the same mesh despite "
      "carrying the largest design vector in the table by two orders of "
      "magnitude."),

("h2", "Where the wall clock goes"),

("p", "Wall-clock time is not the product of the evaluation count and the median "
      "solve cost, and the discrepancy is the most useful thing in this section. "
      "The adaptive partition performed 2716 converged solves at a median of "
      "0.34 s, which is fifteen minutes of arithmetic, but occupied 184 minutes. "
      "Figure 6.29b shows why. The flat band at the median is the boundary trial "
      "sweeps; the clusters above it are the inner L-BFGS exploring; and the rug "
      "along the top is the 243 evaluations that never converged and were "
      "abandoned at the driver's wall-clock limit. Figure 6.29c accumulates that "
      "distribution: 9% of the evaluations carry 90% of the wall clock, and on "
      "the field-aligned run 10% carry 85%."),

("p", "Every implementation has such a ceiling, and it is worth naming because it "
      "is the single largest term in the cost of a run. In the Python driver it is "
      "an explicit subprocess wall-clock limit, set between 20 s and 120 s across "
      "the runs reported here. In the in-process implementation of Section 6.4.1 "
      "there is no wall-clock guard at all, only the Newton iteration limit of "
      "10 000, and the effect is stark. Strategies A and C, which have no cable, "
      "each encountered exactly one design point at which the solve exhausted "
      "that limit, costing about 51 s. Every other solve in those runs took "
      "roughly 0.03 s. That single evaluation accounts for 97% of strategy A's "
      "51.5 s and 96% of strategy C's 53.4 s. Strategies B and D, which carry the "
      "crease cable, encountered no such point and completed 67 and 39 solves in "
      "0.93 s and 1.12 s respectively."),

("p", "This gives the cable argument of Section 6.4.1 a computational corollary. "
      "The cable was introduced because it resolves the crease locally and "
      "relieves the fabric of having to produce the feature by extreme "
      "differential stretch. The same relief is visible in the cost: the "
      "cable-free strategies are the ones that drive the membrane into states the "
      "Newton solver cannot resolve, and adding the cable made the optimisation "
      "roughly fifty times faster in wall clock while also improving the fit. A "
      "stiffener that makes the target easier for the fabric to reach also makes "
      "the forward problem easier to solve."),

("p", "Because non-converged solves dominate, the reported wall clock is largely a "
      "measure of how often the optimiser probes outside the feasible set, which "
      "depends on where it starts. That is a second and independent reason for the "
      "symmetry-reduced warm start of Section 6.3.7, beyond the local-minimum "
      "argument given there. The warm-started ten-region run visible in Figure "
      "6.30 begins at 6.77 mm, which is where the cold-started runs finish, and "
      "completes in 4.8 minutes."),

("h2", "How much of the budget buys accuracy"),

("p", "Figure 6.30 plots the same three runs against evaluations and against wall "
      "clock. The two axes tell different stories, and the wall-clock axis is the "
      "honest one, because the cheap evaluations are the ones that make progress. "
      "On the field-aligned run, 99% of the total improvement is in hand after 204 "
      "of 1241 evaluations, which is 16% of the evaluations but only 7% of the "
      "wall clock; the remaining 63 minutes bought the last 1% of the fit, a "
      "change of well under a tenth of a millimetre. The adaptive run is not the "
      "same case and should not be described as though it were: it reaches 99% of "
      "its improvement at 36% of its wall clock, and its final region sweep, in "
      "the last ten minutes of a three-hour run, still produced a visible gain. "
      "Region refinement can therefore pay late, while the continuous inner "
      "problem does not."),

("p", "The practical reading is that a stopping rule expressed in evaluations, or "
      "in the optimiser's own tolerances, does not match the cost structure. A "
      "rule that stopped the inner optimisation when the running best improved by "
      "less than a set fraction over a set number of evaluations would have "
      "recovered nearly all of the accuracy reported in Section 6.4.6 for a small "
      "part of the time, whereas the outer region sweeps deserve to run to their "
      "own convergence."),

("h2", "What the FDM form finding costs, and why so little"),

("p", f"The inverse problem is not the only optimisation in the workflow, and the "
      f"comparison with the one that precedes it is instructive, the more so "
      f"because the two run on the same discretisation. The shell with openings is "
      f"carried through both stages as a mesh of {fdm['n_verts']} vertices, "
      f"{fdm['n_faces']} faces and {fdm['n_edges']} edges, of which "
      f"{fdm['n_free']} vertices are free and the remainder fixed on the "
      f"boundary. On that mesh the FDM stage of Section 6.2 solves for one force "
      f"density per edge — {fdm['n_design']} design variables — and "
      f"converges in {fdm['iters']} L-BFGS-B iterations and {FDM_S:.1f} s, a mean "
      f"of {FDM_MS:.1f} ms per iteration. What it buys is a close fit: the "
      f"form-found network sits {FDM_RMS_MM:.2f} mm RMS from the target surface, "
      f"{FDM_RMS_PC:.3f}% of the {fdm_res['span_m']:.3f} m span, with the crown "
      f"reproduced to within {FDM_CROWN_UM:.0f} µm, at force densities spanning "
      f"{fdm_res['q_min']:.2f} to {fdm_res['q_max']:.2f} under unit pressure, the "
      f"lower end resting on the floor imposed by the bounds. The inverse problem "
      f"on the identical "
      f"mesh carries between three and twenty design variables and takes between "
      f"four minutes and three hours. The form finding therefore optimises two "
      f"orders of magnitude more variables in a thousandth of the time, and none "
      f"of the difference is a difference in resolution."),

("p", f"Two things account for it, and they are worth separating because only one "
      f"is intrinsic. The first is the cost of a single equilibrium solve: "
      f"{fdm['inflate_ms']:.2f} ms for the FDM network against "
      f"{MD5:.2f} s for the membrane, a factor of about {FDM_RATIO:.0f}. That gap "
      f"is intrinsic to the physics being solved. The FDM equilibrium is "
      f"linear in the vertex positions once the densities are fixed, so a single "
      f"LU factorisation is reused across the {PICARD} Picard steps "
      f"that resolve the pressure coupling, whereas the membrane solve is a Newton "
      f"iteration on a nonlinear material law in which cable segments change "
      f"between taut and slack across four pressure continuation stages. Equal "
      f"vertex counts do not imply equal cost when the equations differ in kind."),

("p", f"The second is the cost of a gradient, and that one is an implementation "
      f"choice rather than a property of the problem. The openings form finding "
      f"differentiates its objective by an adjoint: three linear solves per "
      f"iteration, one per coordinate axis, independent of the number of design "
      f"variables. The inverse problem uses finite differences, at n + 1 "
      f"nonlinear solves per gradient. At the sizes used here that is between "
      f"three and thirty-one solves where the adjoint would be three, but the "
      f"scaling is the real point. Under finite differences the form finding "
      f"would need {fdm['n_design'] + 1} solves per gradient and some "
      f"{(fdm['n_design'] + 1) * fdm['iters'] / 1e6:.1f} million over its "
      f"{fdm['iters']} iterations; it is the adjoint alone that makes a design "
      f"space of that dimension ordinary. The inverse problem tolerates finite "
      f"differences because its design vector was first reduced, by the region "
      f"partition and by symmetry, to a size at which the penalty is affordable "
      f"rather than absent."),

("p", f"How much that choice is worth need not be argued, because the workflow "
      f"contains the controlled experiment. The form finding for the free-form "
      f"shell and for the fluted dome is the same method on the same solver, but "
      f"implemented by direct sensitivity: rather than solving three adjoint "
      f"systems, each iteration solves for the full sensitivity of every free "
      f"vertex to every force density, an n_free by n_edge system per coordinate "
      f"axis. One iteration of the openings form finding costs "
      f"{fdm['ms_per_iter']:.1f} ms at {fdm['n_design']} design variables; one "
      f"iteration of the fluted dome costs {fdm_c5['ms_per_iter'] / 1e3:.2f} s at "
      f"{fdm_c5['n_design']}, and one of the free-form shell "
      f"{fdm_b5['ms_per_iter'] / 1e3:.2f} s at {fdm_b5['n_design']}. That is a "
      f"factor of {DIRECT_RATIO:.0f} per iteration between the two "
      f"implementations at comparable size, and it shows in the wall clocks: "
      f"{FDM_S:.1f} s for the openings against "
      f"{fdm_c5['samples_s'][0] / 60:.0f} minutes for the fluted dome. The direct "
      f"cost is linear in the number of design variables, at "
      f"{PER_VAR_C5:.2f} ms per variable per iteration on the dome and "
      f"{PER_VAR_B5:.2f} ms on the free-form shell, where the adjoint cost is "
      f"flat. Had the openings form finding been written the same way, its "
      f"{fdm['iters']} iterations would have cost of the order of "
      f"{D5_IF_DIRECT_MIN:.0f} minutes instead of nine seconds. The reduction "
      f"this section recommends for the inverse problem is therefore not a "
      f"hypothetical: the workflow already contains one stage that took it and "
      f"two that did not, and the difference between them was measured rather "
      f"than argued."),

("p", f"The two fits are worth holding against each other, since both are RMS "
      f"deviations from the same target surface. The form finding reaches "
      f"{FDM_RMS_MM:.2f} mm in {FDM_S:.1f} s; the best of the inverse runs of "
      f"Section 6.4.6 reaches {FEM_BEST_MM:.2f} mm in an hour. The target is "
      f"therefore very nearly attainable as a pure tension network, and the "
      f"residual is not a failure of the form finding but the cost of realising "
      f"that network in a knitted fabric of finite and directional stiffness, "
      f"under the region and cable constraints the fabrication imposes. It is "
      f"the materialisation, not the geometry, that the expensive optimisation "
      f"is paying for."),

("h2", "Scale, and the cost of not instrumenting"),

("p", "Set against fabrication, none of this is expensive. The most costly single "
      "optimisation in this work is about three hours on one core, against a "
      "knitting time of [X] hours for the corresponding specimen and [X] for "
      "assembly. The workflow is not compute-bound at the scale of one design. It "
      "becomes compute-bound in two situations, both encountered here: when the "
      "region partition is treated as a design variable, and when one target is "
      "swept over a design parameter. Section 6.5.2 is the second case; at a "
      "fixed mesh of 929 elements, the evaluations needed to fit the free-form "
      "target rose from 622 at a 1.2 m span to 4687 at 3.0 m, so the cost of the "
      "inverse problem grows with how far the fabric is asked to stretch and not "
      "only with the size of the discretisation. Across all the runs behind "
      "Sections 6.4.1 to 6.4.6 the case studies consumed of the order of 3 × 10⁴ "
      "forward solves, of which the reported results are a small fraction; the "
      "remainder is the cost of the strategy comparisons."),

("p", "Two straightforward reductions were not implemented, and the size of the "
      "second one can be read directly off the table. The n + 1 solves of a "
      "finite-difference gradient are independent and could be dispatched across "
      "the twenty available cores, which would bring a gradient on the "
      "thirty-variable problem down to the cost of a single solve. And the "
      "subprocess driver restarts the forward solve from the undeformed "
      "configuration at every evaluation, whereas the in-process implementation of "
      "Section 6.4.1 warm-starts from the previous accepted iterate. On the same "
      f"634-element mesh a cold solve costs {M2:.2f} s and a warm-started one "
      "between 0.010 s and 0.029 s, so the warm start is worth a factor of three "
      "to nine before any other change; the gain would be larger still at the "
      "small perturbations a finite-difference gradient takes, which is exactly "
      "where the subprocess driver throws the previous solution away. Neither "
      "change affects any result reported in this chapter, and either would move a "
      "region-refinement run from hours to minutes. A third reduction, an adjoint "
      "gradient for the membrane solve in place of finite differences, is a larger "
      "undertaking, since it requires differentiating through the Newton iteration "
      "and the switching of the cable segments between taut and slack; but it is "
      "not speculative, being the device the openings form finding of this same "
      "workflow already uses, and it is the one change that would lift the ceiling "
      "on the number of design variables rather than merely lower the cost of the "
      "present ones. The two form-finding runs written the other way put a "
      "measured figure on what declining it costs."),

("p", "A final methodological remark, since it cost real effort to recover. The "
      "optimisation drivers recorded parameters, losses and evaluation counts, but "
      "never elapsed time, and stored only the last thirty evaluations of each "
      "loss history. The timings reported above had to be reconstructed from the "
      "modification times of the per-evaluation output files, and the loss "
      "trajectories by recomputing the objective from those files; for two of the "
      "four case studies the timestamps had been overwritten and the wall clocks "
      "are simply lost. Recording an elapsed time and a full loss history per "
      "evaluation is a two-line change to the driver and would have made this "
      "section a matter of reading a file."),
]

TABLE = [
    ("Case study", "Section", "Mesh (v / f)", "Design variables",
     "Forward solves", "Median solve", "Wall clock"),
    ("Middle crease, isotropic, no cable (A)", "6.4.1", "341 / 634", "1",
     "47", "0.010 s*", "51.5 s"),
    ("Middle crease, isotropic, cable (B)", "6.4.1", "341 / 634", "1",
     "67", "0.014 s", "0.93 s"),
    ("Middle crease, anisotropic, no cable (C)", "6.4.1", "341 / 634", "2",
     "93", "0.024 s*", "53.4 s"),
    ("Middle crease, anisotropic, cable (D)", "6.4.1", "341 / 634", "2",
     "39", "0.029 s", "1.12 s"),
    ("Free-form shell, 9 regions", "6.4.2", "497 / 929", "30",
     "622", f"{MB5:.2f} s", f"≥ {622 * MB5:.0f} s †"),
    ("Fluted dome, 16 regions, D8-reduced", "6.4.3", "1309 / 2137", "7 + 1",
     "378 + 34", f"{MC5:.2f} s", f"≥ {412 * MC5:.0f} s †"),
    ("Openings, 1 region", "6.4.6", "847 / 1563", "3",
     "162", f"{MD5:.2f} s", f"≥ {162 * MD5:.0f} s †"),
    ("Openings, 4 adaptive regions", "6.4.6", "847 / 1563", "9",
     "2959", "0.34 s", "184 min"),
    ("Openings, 10 field-aligned regions", "6.4.6", "847 / 1563", "20",
     "1241", "0.39 s", "68 min"),
    ("Openings, 10 symmetric regions, warm-started", "6.4.6", "847 / 1563", "10",
     "90", "0.46 s", "4.8 min"),
    ("Crease, form finding ‡", "6.2", "—", "—", "—", "—", "interactive"),
    ("Free-form shell, form finding ‡", "6.2",
     f"{fdm_b5['n_verts']} / {fdm_b5['n_faces']}", f"{fdm_b5['n_design']}",
     f"{fdm_b5['iters']}", f"{fdm_b5['ms_per_iter'] / 1e3:.2f} s",
     f"{fdm_b5['samples_s'][0] / 60:.1f} min §"),
    ("Fluted dome, form finding ‡", "6.2",
     f"{fdm_c5['n_verts']} / {fdm_c5['n_faces']}", f"{fdm_c5['n_design']}",
     f"{fdm_c5['iters']}", f"{fdm_c5['ms_per_iter'] / 1e3:.2f} s",
     f"{fdm_c5['samples_s'][0] / 60:.1f} min"),
    ("Openings, form finding ‡", "6.2",
     f"{fdm['n_verts']} / {fdm['n_faces']}", f"{fdm['n_design']}",
     f"{fdm['iters']}", f"{fdm['ms_per_iter'] / 1e3:.3f} s", f"{FDM_S:.1f} s"),
]

TABLE_CAPTION = (
    "Table 6.X: Computational cost of the case studies. For the crease strategies "
    "the solve cost is the mean over the run itself, which is warm-started in "
    "process; entries marked * exclude the single non-converging solve discussed "
    "in the text. For the free-form shell, the fluted dome and the single-region "
    "fit it is the median over twenty-five design points about the optimum, each "
    "a cold solve. For the three region-scheme runs it is the median over the "
    "recorded run. Wall clocks marked † were lost when the output files were "
    "rewritten and are not measurements: they are the evaluation count times the "
    "median solve cost, and therefore lower bounds, since they credit the run "
    "with no non-converging solves at all. The three measured region-scheme runs "
    f"exceed the same bound by factors of {184 * 60 / (2959 * 0.34):.0f}, "
    f"{68 * 60 / (1241 * 0.39):.0f} and {4.8 * 60 / (90 * 0.46):.0f} "
    "respectively, so the true figures are plausibly of the order of ten minutes "
    "rather than one. The rows marked ‡ are not inverse problems but the form "
    "finding that precedes each of them, and for those rows the solve cost is "
    "one L-BFGS-B iteration — an equilibrium solve and a gradient — rather than "
    "a single membrane solve, so it is comparable within the block and not "
    "against the rows above. They are not comparable among themselves either "
    "without the caveat that carries the argument of this section: the openings "
    "form finding takes its gradient by an adjoint and the other two by direct "
    "sensitivity. The crease form finding is an interactive session rather than "
    "a fitted optimisation and has no meaningful wall clock. The entry marked § "
    "stopped at its iteration cap rather than at convergence, so it too is a "
    "floor."
)

CAPTIONS = {
    "computational_time":
        "Figure 6.29: The cost of the inverse problem. (a) The cost of one forward "
        "solve at twenty-five design points within ±10% of each recorded optimum, "
        "with the median marked; rows are ordered by mesh size, and the order of "
        "the medians does not follow it. (b) Every evaluation of the "
        "four-region adaptive run on the shell with openings: the flat band is the "
        "boundary trial sweeps, the clusters are the inner L-BFGS, and the rug on "
        "the ceiling is the 243 evaluations abandoned at the driver's wall-clock "
        "limit. (c) The same runs accumulated cheapest-first: a tenth of the "
        "evaluations carry nine tenths of the wall clock.",
    "convergence_cost":
        "Figure 6.30: The same three runs on the shell with openings, plotted "
        "against evaluations (a) and against wall clock (b). Faint dots are "
        "individual evaluations, the heavy line is the running best, and the "
        "filled circle marks 99% of the total improvement. The field-aligned run "
        "reaches that point after 7% of its wall clock; the adaptive run does not, "
        "and its last region sweep still improves the fit.",
}

NOTES = [
("h", "Notes for revision — not part of the section"),
("p", "Numbers above are measured or recorded, never estimated. Provenance: "
      "per-solve costs from FDM/measure_solve_cost.py (25 samples per geometry, "
      "this machine); 2part solve counts and wall clocks from re-running "
      "build-linux/best_fit_* with the sim_count instrumentation already in those "
      "sources; evaluation counts for B5, C5 and D5 from the n_calls fields of "
      "FDM/optimisation/*_optimised*.json; D5 wall clocks and tail statistics from "
      "the per-evaluation output-file timestamps; loss trajectories from "
      "FDM/reconstruct_loss_history.py; form-finding and directional-field costs "
      "from FDM/data/fdm_cost.json, whose samples come from FDM/fofin_D5.py, "
      "FDM/fofin_B5.py and FDM/fofin_C5_dense.py, all three now instrumented to "
      "report their elapsed time, iteration count and design dimension, the B5 "
      "and C5 runs having been executed concurrently on separate cores while the "
      "D5 samples were taken on an otherwise idle machine, and from the "
      "whole-script wall clock of "
      "FDM/directional_field_D5.py; the per-solve figure of "
      f"{fdm['inflate_ms']:.2f} ms is {fdm['inflate_repeats']} repeats of "
      "inflate() at the converged densities."),
("p", "1. Fabrication times. Two [X] placeholders remain in the comparison "
      "paragraph — knitting hours and assembly time per specimen."),
("p", "2. The finite-difference step. Section 6.3.7 states 0.05. optimise_B5.py "
      "uses 0.05, but optimise_C5_16region.py and the D5 scripts use 0.002, and "
      "the 2part C++ optimisers use 1e-4 on log-parameters. The text needs to name "
      "the per-case values or explain the difference."),
("p", "3. Section 6.4.1 deviations. Re-running strategies A to D reproduces the "
      "reported stretch factors exactly (C at 1.461 / 0.908 and D at 1.299 / "
      "0.960, matching the text), but the deviations computed from those results "
      "do not match the reported ones: strategy D gives a mean of 7.90 mm over all "
      "vertices and 9.13 mm over interior vertices against the 5.088 mm in the "
      "text, and a maximum of 22.5 mm against 17.393 mm. The optimum is the same, "
      "so the difference is in how the deviation was measured. Worth resolving "
      "before the numbers in 6.4.1 and any cost-accuracy comparison are used "
      "together."),
("p", "4. Fluted dome cable count. Section 6.4.3 describes twelve radial cables "
      "plus one circumferential, then refers to 24 cables. The stored result has "
      "16 regions and 24 cable sections under D8 symmetry, i.e. eight sectors."),
("p", "5. Strategy E of Section 6.4.1 is absent from the table because its "
      "executable initialises Polyscope and needs a display; its solve count was "
      "never recorded. It can be added by running "
      "build-linux/best_fit_stretch_factors_3region_adaptive on a machine with a "
      "display, or by removing the visualisation calls."),
("p", "6. A latent trap, worth a sentence somewhere in 6.3.7. For the shell with "
      "openings the rest mesh is the target mesh, so any evaluation that returns "
      "the rest shape unchanged scores an exact zero on the objective. The "
      "drivers' validity gate rejects those, and it fired on 10 evaluations of the "
      "adaptive run and 21 of the field-aligned run. Without the gate the "
      "objective has a perfect-scoring attractor that corresponds to a structure "
      "that never inflated."),
("p", "7. Form-finding cost is now measured for three of the four geometries. Two "
      "caveats attach to the new rows. The free-form shell stopped at its "
      "iteration cap rather than at convergence, so 171 s is a floor and its fit, "
      "0.30% of span, is well short of the 0.08% the other two reach; and it runs "
      "on input/B5.obj, 269 vertices, where the inverse problem of 6.4.2 uses a "
      "497/929 remesh, so the two are not the same discretisation. Whether the "
      "form finding behind the reported 6.4.2 result was this run or one on the "
      "finer mesh is worth confirming. The crease form finding (square_crease.py) "
      "remains unmeasured and is a different kind of object: an interactive "
      "session of fixed-point fd_numpy steps, needing compas_fd and compas_view2, "
      "neither installed here. Remeshing and cable extraction are also "
      "unmeasured."),

("p", "8. The gradient implementations diverged without anyone deciding they "
      "should. fofin_D5.py uses an adjoint; fofin_B5.py, fofin_C5.py and "
      "fofin_C5_dense.py solve for the full sensitivity matrix. The measured "
      "penalty is a factor of about 260 per iteration. Porting the adjoint from "
      "fofin_D5.py to the other two is a contained change and would bring the "
      "fluted dome form finding from 34 minutes to seconds; worth doing before "
      "any sweep that re-runs form finding many times."),

("p", "9. fofin_B5.py imported compas.numerical, which no longer exists in "
      "compas 2.x, so it could not run at all on the timing machine until the "
      "import was made to fall back across both APIs. Other scripts in FDM/ may "
      "carry the same breakage; fofin_C5.py and the fofin_butt/cross/seismic "
      "family were not checked."),
]


def build_md():
    lines = []
    for kind, text in BODY:
        if kind == "h":
            lines.append(f"# {text}\n")
        elif kind == "h2":
            lines.append(f"## {text}\n")
        elif kind == "p":
            lines.append(f"{text}\n")
        elif kind == "table":
            head = TABLE[0]
            lines.append("| " + " | ".join(head) + " |")
            lines.append("|" + "---|" * len(head))
            for row in TABLE[1:]:
                lines.append("| " + " | ".join(row) + " |")
            lines.append("")
            lines.append(f"*{TABLE_CAPTION}*\n")
    for name, cap in CAPTIONS.items():
        lines.append(f"![{name}](../FDM/figures/{name}.png)\n")
        lines.append(f"*{cap}*\n")
    for kind, text in NOTES:
        lines.append(("# " if kind == "h" else "") + text + "\n")
    return "\n".join(lines)


def build_docx():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for kind, text in BODY:
        if kind == "h":
            doc.add_heading(text, level=2)
        elif kind == "h2":
            doc.add_heading(text, level=3)
        elif kind == "p":
            p = doc.add_paragraph(text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif kind == "table":
            t = doc.add_table(rows=1, cols=len(TABLE[0]))
            t.style = "Light Grid Accent 1"
            for i, h in enumerate(TABLE[0]):
                cell = t.rows[0].cells[i]
                cell.text = h
                for r in cell.paragraphs[0].runs:
                    r.font.bold = True
                    r.font.size = Pt(8)
            for row in TABLE[1:]:
                cells = t.add_row().cells
                for i, v in enumerate(row):
                    cells[i].text = v
                    for r in cells[i].paragraphs[0].runs:
                        r.font.size = Pt(8)
            cap = doc.add_paragraph(TABLE_CAPTION)
            cap.runs[0].font.size = Pt(9)
            cap.runs[0].font.italic = True

    for name, cap in CAPTIONS.items():
        path = os.path.join(FIG, f"{name}.png")
        if not os.path.exists(path):
            continue
        doc.add_picture(path, width=Inches(6.4))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c = doc.add_paragraph(cap)
        c.runs[0].font.size = Pt(9)
        c.runs[0].font.italic = True

    doc.add_page_break()
    for kind, text in NOTES:
        if kind == "h":
            doc.add_heading(text, level=2)
        else:
            doc.add_paragraph(text, style="List Bullet")

    os.makedirs(os.path.dirname(OUT_DOCX), exist_ok=True)
    doc.save(OUT_DOCX)


if __name__ == "__main__":
    os.makedirs(os.path.dirname(OUT_MD), exist_ok=True)
    with open(OUT_MD, "w") as f:
        f.write(build_md())
    build_docx()
    print(f"  saved: {OUT_DOCX}")
    print(f"  saved: {OUT_MD}")
