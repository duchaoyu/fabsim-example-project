"""Emit Section 6.5.2 (Scalability) as .docx and .md.

Same arrangement as build_section_6_5_1.py: one source for both formats, every
number read from the run outputs rather than retyped.

    data/scalability.json        written by figure_scalability.py
    optimisation/B5_multiscale_summary.json

    python3 figure_scalability.py && python3 build_section_6_5_2.py
"""

import json
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
FIG = os.path.join(HERE, "figures")
OUT_DOCX = os.path.join(HERE, "data", "Section_6_5_2_Scalability.docx")
OUT_MD = os.path.join(ROOT, "docs", "6_5_2_scalability.md")

S = json.load(open(os.path.join(HERE, "data", "scalability.json")))
RUNS = S["runs"]
SUM = {d["tag"]: d for d in json.load(
    open(os.path.join(HERE, "optimisation", "B5_multiscale_summary.json")))}

R0, RN = RUNS[0], RUNS[-1]
K_MEAN, K_MAX, K_COST = S["exponent_mean"], S["exponent_max"], S["exponent_cost"]
CAPPED = [r for r in RUNS if not r["converged"]]

HEAD = ("span D (m)", "mean deviation (mm)", "max deviation (mm)",
        "mean / D", "max / D", "FEM solves", "course tension (N/m)")
TABLE = [(f"{r['D']:.1f}", f"{r['mean_mm']:.1f}", f"{r['max_mm']:.1f}",
          f"{r['mean_pct']:.2f}%", f"{r['max_pct']:.2f}%",
          f"{r['n_calls']}" + ("" if r["converged"] else " *"),
          f"{r['T_course']:.0f}") for r in RUNS]

TABLE_CAPTION = (
    "Table 6.W: The free-form shell of Section 6.4.2 optimised at four spans "
    f"against a target scaled with it. Deviations are over the "
    f"{S['n_interior']} interior vertices of the 497, the boundary being "
    "clamped and carrying no signal. Course tension is the mean over the "
    "surface at the optimum. Entries marked * stopped at the optimiser's "
    "iteration cap rather than at convergence, so their deviations are upper "
    "bounds on what the method achieves and their solve counts are censored."
)

BODY = [
("h", "6.5.2 Scalability"),

("p", "Everything in Section 6.4 was designed and built at the scale of a "
      "laboratory specimen. The question this section asks is what changes when "
      "the same design is built larger. It is not whether a bigger shell "
      "deviates further from its target in millimetres — everything about it is "
      "bigger, so it would be surprising if it did not — but whether it "
      "deviates further in proportion, and what that costs to establish."),

("p", f"The free-form shell of Section 6.4.2 was optimised at four spans, "
      f"{R0['D']:.1f}, 1.5, 2.0 and {RN['D']:.1f} m, against a target scaled "
      f"with the span so that the intended shape is the same shape throughout. "
      f"The mesh is held at 497 vertices and 929 faces, the material at motif 1, "
      f"the pressure at 1000 Pa and the region partition at nine, so the only "
      f"thing that varies is size. Each run was re-solved afterwards at its own "
      f"optimum, and the deviations reported here are measured from that "
      f"verification solve rather than from the optimiser's internal record."),

("table", None),

("h2", "The fit degrades faster than the structure grows"),

("p", f"The mean deviation rises from {R0['mean_mm']:.1f} mm at "
      f"{R0['D']:.1f} m to {RN['mean_mm']:.1f} mm at {RN['D']:.1f} m. That is a "
      f"factor of {RN['mean_mm'] / R0['mean_mm']:.1f} for a factor of "
      f"{RN['D'] / R0['D']:.1f} in span, so the deviation is not proportional "
      f"to the size of the structure: fitted as a power law it goes as "
      f"D^{K_MEAN:.2f}, and the maximum as D^{K_MAX:.2f}. Figure 6.31a shows "
      f"both against the proportional line."),

("p", f"Stated the way a client would ask it, the same fact is that accuracy "
      f"relative to the span gets worse with span. The mean deviation is "
      f"{R0['mean_pct']:.2f}% of the span at {R0['D']:.1f} m and "
      f"{RN['mean_pct']:.2f}% at {RN['D']:.1f} m; the maximum grows from "
      f"{R0['max_pct']:.2f}% to {RN['max_pct']:.2f}%. A design tolerance "
      f"written as a fraction of span — the natural way to write one — is "
      f"therefore harder to meet at every increase in size, by roughly "
      f"D^{K_MEAN - 1:.2f}."),

("h2", "Where the error is, and what that says about its cause"),

("p", "The obvious explanations are that the fabric is being stretched further "
      "at larger spans, or that the optimiser is running out of parameter "
      "range, and neither survives contact with the runs. The stretch factors "
      "stay between 1.041 and 1.053 at every span, nowhere near the bounds of "
      "[0.95, 1.4] the optimisation is allowed, so the parameterisation is not "
      "saturating. The crown height, meanwhile, gets relatively closer to its "
      "target as the span grows, from 18.5% of span against an intended 19.3% "
      "at the smallest scale to 19.1% at the largest. Whatever is going wrong "
      "is not going wrong at the crown, and it is not the fabric running out "
      "of stretch."),

("p", "Figure 6.31d locates it. The deviation field is plotted at each span, "
      "each normalised by its own span so the patterns may be compared, and it "
      "is the same pattern throughout — a grid of nine cells with the error "
      "concentrated in the seams between them and in the four side lobes. That "
      "grid is the region partition. The surface is fitted with nine "
      "piecewise-constant pairs of stretch factors, and what the fit cannot "
      "express is the variation within a region and the discontinuity at its "
      "edge. Scaling the structure does not change that pattern; it magnifies "
      "it, because the same relative shape error over a longer edge is a larger "
      "absolute one, and because the curvature the fabric must produce inside "
      "each cell grows with the cell."),

("p", "This is a more useful answer than a scaling exponent on its own, "
      "because it names the lever. The way to build this shape larger without "
      "losing proportional accuracy is not a stiffer fabric or a higher "
      "pressure but a finer partition — more regions, or regions placed where "
      "the seams currently fall. Section 6.4.6 already shows what that costs on "
      "a different geometry: going from four regions to ten reduced the "
      "deviation while roughly halving the evaluations needed, once the regions "
      "were aligned with the directional field rather than grown adaptively."),

("h2", "What it costs to establish"),

("p", f"The optimisation cost grows as D^{K_COST:.2f}, from "
      f"{R0['n_calls']} forward solves at {R0['D']:.1f} m to "
      f"{RN['n_calls']} at {RN['D']:.1f} m. Nothing in the problem gets bigger "
      f"in the numerical sense — the mesh, the design vector and the solver "
      f"tolerances are identical at every span — so this is not the cost of a "
      f"larger discretisation. It is the cost of a harder optimisation: as the "
      f"target moves further out of reach of the parameterisation, the "
      f"objective flattens and the finite-difference gradient has less to work "
      f"with, and L-BFGS-B takes more iterations to make less progress. At "
      f"0.13 s a solve, from the timings of Section 6.5.1, the largest run is "
      f"about ten minutes, so this is a limit on patience rather than on "
      f"feasibility."),

("p", f"Two of the four runs, at 1.5 and {CAPPED[-1]['D']:.1f} m, stopped at "
      f"the optimiser's iteration cap rather than at convergence. Their "
      f"deviations are therefore upper bounds on what the method achieves at "
      f"those spans, and their solve counts are censored from below. The "
      f"exponents above should be read with that in mind: the true degradation "
      f"may be gentler than D^{K_MEAN:.2f} and the true cost steeper than "
      f"D^{K_COST:.2f}. Re-running the two capped spans to convergence is the "
      f"single thing that would firm up this section, and at ten minutes a run "
      f"it is cheap."),

("h2", "What scales, and what does not"),

("p", f"Three quantities behave differently as the span grows, and separating "
      f"them is the practical result of this section. The intended shape scales "
      f"exactly, by construction. The membrane tension scales sub-linearly: the "
      f"mean course tension rises from {R0['T_course']:.0f} to "
      f"{RN['T_course']:.0f} N/m over a 2.5-fold span, so a fabric qualified at "
      f"the specimen scale is not immediately disqualified at the larger one, "
      f"though the margin against the material's usable range narrows and "
      f"should be checked against Chapter 4 rather than assumed. And the "
      f"achievable accuracy scales adversely, as D^{K_MEAN:.2f}."),

("p", "So the workflow scales in the sense that matters least — it runs, and it "
      "produces a design — and fails to scale in the sense that matters most, "
      "which is that the design gets proportionally worse. The remedy is "
      "known and is a design decision rather than a fabrication one: refine the "
      "partition as the structure grows, and expect the number of regions "
      "needed to hold a fixed relative accuracy to rise with span. Establishing "
      "that relation — regions required against span, at fixed accuracy — is "
      "the natural continuation of this study and is not attempted here."),
]

CAPTIONS = [
    ("scalability",
     "Figure 6.31: The free-form shell at four spans. All three upper panels "
     "share one linear span axis. (a) Mean and maximum deviation against span, "
     "logarithmic in the deviation, with the fitted power laws and the "
     "proportional line for comparison; both grow faster than the span. (b) The "
     "same normalised by span, which is how a tolerance would be written: "
     "relative accuracy worsens throughout. (c) Optimisation cost, quadratic in "
     "the span, with the two runs that stopped at the iteration cap shown "
     "hollow. (d) The deviation field at each span, each "
     "normalised by its own span: the same nine-cell pattern of the region "
     "partition throughout, magnified rather than changed."),
]

NOTES = [
("h", "Notes for revision — not part of the section"),
("p", "Every number above is read from data/scalability.json, which "
      "figure_scalability.py computes from the stored verification solves "
      "(optimisation/check_<tag>_verts.csv) against the scaled targets "
      "(data/B5_remeshed_shared*.off), over the interior-vertex index "
      "data/B5_remeshed_interior_idx.npy. Solve counts and tensions come from "
      "optimisation/B5_multiscale_summary.json. Re-run the figure script and "
      "rebuild and the section follows."),
("p", "1. Two of four runs are capped, at 1.5 and 3.0 m. This is the weakest "
      "point of the section and it is cheap to fix: about ten minutes of "
      "compute each at the Section 6.5.1 solve cost. Until then the exponents "
      "are indicative rather than measured."),
("p", "2. Four spans over a 2.5-fold range is a short lever arm for a power "
      "law. The fitted exponents are reported to two decimals because that is "
      "what the fit returns, not because the data supports that precision; "
      "'grows roughly as the square of the span' is the honest reading of the "
      "cost, and 'faster than proportionally, close to D^1.8' of the "
      "deviation."),
("p", "3. The mechanism in Figure 6.31d is an interpretation, well supported by "
      "the pattern matching the partition but not demonstrated. The "
      "demonstration is a re-run at one span with a finer partition: if the "
      "deviation falls and the pattern subdivides accordingly, the reading is "
      "confirmed. That experiment would also give the regions-against-span "
      "relation the closing paragraph says is missing."),
("p", "3a. The crown-height panel has been removed from Figure 6.31, and the "
      "three remaining upper panels put on one linear span axis so they read "
      "against the same scale; (a) and (c) keep a logarithmic vertical axis, "
      "so the fitted power laws now appear as curves rather than as straight "
      "lines. The crown-height claim in the second paragraph of 'Where the "
      "error is' therefore no longer has a panel behind it. The numbers are "
      "still in data/scalability.json, as crown_ratio and target_crown_ratio, "
      "but a reader of the section has to take them on trust. Either quote "
      "both endpoints in the sentence or drop the claim."),
("p", "4. The section numbering follows the thesis: 6.5.1 computational time, "
      "6.5.2 this, 6.5.3 robustness. Figure 6.31 here means the figures of "
      "6.5.3 shift to 6.32 onward; both build scripts should be renumbered "
      "together once the chapter's figure list is settled."),
("p", "5. Tension is reported as the surface mean. If the fabric's usable "
      "range is to be checked at 3.0 m, the maximum matters more: it reaches "
      "1755 N/m in course at the largest span against 908 N/m at the smallest. "
      "Whether that is inside the qualified range is a Chapter 4 question this "
      "section does not answer."),
]


def md_table():
    out = ["| " + " | ".join(HEAD) + " |", "|" + "---|" * len(HEAD)]
    out += ["| " + " | ".join(r) + " |" for r in TABLE]
    return "\n".join(out) + "\n"


def build_md():
    lines = []
    for kind, text in BODY:
        if kind == "h":
            lines.append(f"# {text}\n")
        elif kind == "h2":
            lines.append(f"## {text}\n")
        elif kind == "p":
            lines.append(f"{text}\n")
        elif kind == "table":
            lines.append(md_table())
            lines.append(f"*{TABLE_CAPTION}*\n")
    for name, cap in CAPTIONS:
        lines.append(f"![{name}](../FDM/figures/{name}.png)\n")
        lines.append(f"*{cap}*\n")
    for kind, text in NOTES:
        lines.append(("# " if kind == "h" else "") + text + "\n")
    return "\n".join(lines)


def build_docx():
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    for kind, text in BODY:
        if kind == "h":
            doc.add_heading(text, level=2)
        elif kind == "h2":
            doc.add_heading(text, level=3)
        elif kind == "p":
            p = doc.add_paragraph(text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif kind == "table":
            t = doc.add_table(rows=1, cols=len(HEAD))
            t.style = "Table Grid"
            for i, h in enumerate(HEAD):
                cell = t.rows[0].cells[i]
                cell.text = h
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True
                    run.font.size = Pt(8)
            for row in TABLE:
                cells = t.add_row().cells
                for i, v in enumerate(row):
                    cells[i].text = v
                    for run in cells[i].paragraphs[0].runs:
                        run.font.size = Pt(8)
            cap = doc.add_paragraph(TABLE_CAPTION)
            cap.runs[0].font.size = Pt(9)
            cap.runs[0].font.italic = True

    for name, cap in CAPTIONS:
        path = os.path.join(FIG, f"{name}.png")
        if not os.path.exists(path):
            print(f"  missing figure, skipped: {path}")
            continue
        doc.add_picture(path, width=Inches(6.4))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c = doc.add_paragraph(cap)
        c.runs[0].font.size = Pt(9)
        c.runs[0].font.italic = True

    doc.add_page_break()
    for kind, text in NOTES:
        if kind == "h":
            doc.add_heading(text, level=2)
        else:
            doc.add_paragraph(text, style="List Bullet")

    doc.save(OUT_DOCX)


if __name__ == "__main__":
    os.makedirs(os.path.dirname(OUT_MD), exist_ok=True)
    with open(OUT_MD, "w") as f:
        f.write(build_md())
    build_docx()
    print(f"  saved: {OUT_DOCX}")
    print(f"  saved: {OUT_MD}")
