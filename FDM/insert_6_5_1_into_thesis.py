"""Insert Section 6.5.1 into the thesis document, replacing whatever is there.

The section itself is authored in build_section_6_5_1.py, which owns the text and
pulls every number from the measurement files. This script only places it: it
finds the existing "6.5.1 Computational time" heading in the thesis, deletes the
body between that heading and the next one of the same level, and writes the
current section in its place, using the thesis's own styles.

It is idempotent. Run it again after editing build_section_6_5_1.py and the
section in the thesis is replaced, not duplicated. The revision notes are never
inserted; they belong to the working copy only.

    python3 insert_6_5_1_into_thesis.py [--dry-run]

The thesis is not under version control, so a timestamped backup is written
beside it before anything is saved.
"""

import os
import shutil
import sys
from datetime import datetime

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx import Document

import build_section_6_5_1 as sec

HERE = os.path.dirname(os.path.abspath(__file__))
THESIS = os.path.join(HERE, "data", "Knitting for Shell structures.docx")
FIG = os.path.join(HERE, "figures")

HEADING_TEXT = "6.5.1 Computational time"
HEADING_LEVEL = "Heading 3"
SUBHEAD_LEVEL = "Heading 4"
BODY_STYLE = "Normal"
PIC_WIDTH = Inches(6.2)


def find_section(doc):
    """Return (heading paragraph, first paragraph after the section)."""
    paras = doc.paragraphs
    start = None
    for i, p in enumerate(paras):
        if p.style.name == HEADING_LEVEL and p.text.strip().replace(
                " ", "").startswith("6.5.1"):
            start = i
            break
    if start is None:
        raise SystemExit("6.5.1 heading not found in the thesis")

    for j in range(start + 1, len(paras)):
        if paras[j].style.name in (HEADING_LEVEL, "Heading 2", "Heading 1"):
            return paras[start], paras[j]
    return paras[start], None


def clear_between(heading, stop):
    """Delete every block element between the heading and the next section."""
    body = heading._p.getparent()
    stop_el = stop._p if stop is not None else None
    removed = 0
    el = heading._p.getnext()
    while el is not None and el is not stop_el:
        nxt = el.getnext()
        body.remove(el)
        removed += 1
        el = nxt
    return removed


class Inserter:
    """Appends blocks at the end of the document, then moves them into place."""

    def __init__(self, doc, anchor):
        self.doc = doc
        self.anchor = anchor._p

    def _move(self, element):
        self.anchor.addnext(element)
        self.anchor = element

    def paragraph(self, text, style=BODY_STYLE, size=None, italic=False,
                  justify=True, center=False):
        p = self.doc.add_paragraph(text, style=style)
        if justify:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            if size:
                r.font.size = Pt(size)
            if italic:
                r.font.italic = True
        self._move(p._p)
        return p

    def heading(self, text):
        p = self.doc.add_paragraph(text, style=SUBHEAD_LEVEL)
        self._move(p._p)
        return p

    def table(self, rows):
        t = self.doc.add_table(rows=1, cols=len(rows[0]))
        t.style = "Table Grid"
        for i, h in enumerate(rows[0]):
            cell = t.rows[0].cells[i]
            cell.text = h
            for r in cell.paragraphs[0].runs:
                r.font.bold = True
                r.font.size = Pt(8)
        for row in rows[1:]:
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = v
                for r in cells[i].paragraphs[0].runs:
                    r.font.size = Pt(8)
        self._move(t._tbl)
        return t

    def picture(self, path):
        self.doc.add_picture(path, width=PIC_WIDTH)
        p = self.doc.paragraphs[-1]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._move(p._p)
        return p


def main():
    dry = "--dry-run" in sys.argv
    doc = Document(THESIS)
    heading, stop = find_section(doc)

    print(f"  found:   {heading.text.strip()!r} ({heading.style.name})")
    print(f"  next:    {stop.text.strip()!r}" if stop is not None
          else "  next:    end of document")
    removed = clear_between(heading, stop)
    print(f"  cleared: {removed} block(s) of existing content")

    heading.text = HEADING_TEXT
    heading.style = doc.styles[HEADING_LEVEL]

    ins = Inserter(doc, heading)
    for kind, text in sec.BODY:
        if kind == "h":
            continue                      # the thesis already numbers the section
        elif kind == "h2":
            ins.heading(text)
        elif kind == "p":
            ins.paragraph(text)
        elif kind == "table":
            ins.table(sec.TABLE)
            ins.paragraph(sec.TABLE_CAPTION, style="Caption", size=9,
                          italic=True, justify=False)

    for name, cap in sec.CAPTIONS.items():
        path = os.path.join(FIG, f"{name}.png")
        if not os.path.exists(path):
            print(f"  missing figure, skipped: {path}")
            continue
        ins.picture(path)
        ins.paragraph(cap, style="Caption", size=9, italic=True, justify=False)

    if dry:
        print("  dry run, nothing written")
        return

    stamp = datetime.now().strftime("%Y%m%d%H%M%S")
    backup = THESIS.replace(".docx", f".backup_{stamp}.docx")
    shutil.copy2(THESIS, backup)
    print(f"  backup:  {backup}")
    doc.save(THESIS)
    print(f"  saved:   {THESIS}")


if __name__ == "__main__":
    main()
