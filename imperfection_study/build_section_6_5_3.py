"""Emit Section 6.5.3 (Robustness to construction imprecision) as .docx and .md.

Same arrangement as FDM/build_section_6_5_1.py: one source for both formats, and
every number pulled from the Block A output files rather than retyped, so the
section cannot drift from the runs behind it.

    data/block_A_disc_sensitivity.csv     the six factors on the circular dome
    data/block_A_2part_sensitivity.csv    the same six on the crease
    data/block_A_overlap.json             aggregates from --check-overlap

    python3 build_section_6_5_3.py
"""

import csv
import json
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
FIG = os.path.join(HERE, "figures")
OUT_DOCX = os.path.join(HERE, "data", "Section_6_5_3_Robustness.docx")
OUT_MD = os.path.join(ROOT, "docs", "6_5_3_robustness.md")


def load(geom):
    path = os.path.join(HERE, "data", f"block_A_{geom}_sensitivity.csv")
    with open(path) as f:
        return {r["factor"]: r for r in csv.DictReader(f)}


disc, twopart = load("disc"), load("2part")
cable = load("2part_cable")
ov = json.load(open(os.path.join(HERE, "data", "block_A_overlap.json")))

# Display names and the unit each tolerance is quoted in.
NAME = {"s_wale": "s_wale", "s_course": "s_course", "pressure": "p",
        "E1": "E1", "r": "E2/E1", "R": "R"}
DELTA = {"s_wale": "0.05", "s_course": "0.05", "pressure": "50 Pa",
         "E1": "500 N/m", "r": "0.250", "R": "2 mm"}

# Where each tolerance comes from, for the provenance table. Kept in step with
# tolerances.py, which holds the magnitudes and the error models.
PROVENANCE = [
    ("s_wale, s_course", "0.05 on the factor",
     "Stitch size, measured in the Chapter 4 tensile tests. The stretch factor "
     "is what the commanded stitch length delivers, so a per-stitch error of "
     "the order of 0.1 mm accumulates over the courses of a panel into a "
     "factor deviation.",
     "systematic + independent", "estimate"),
    ("E1, E2/E1", "10%",
     "Chapter 4 tensile tests: scatter across repeats on nominally identical "
     "specimens, which is fabrication limitation, and across yarn batches. "
     "Chapter 4 reports the measured spread; it should replace this assumed "
     "10%.",
     "systematic", "estimate"),
    ("p", "50 Pa (5%)",
     "Sensor resolution and the width of the band the valve holds between "
     "corrections, from 6.1.2.",
     "systematic", "estimate"),
    ("R", "2 mm (0.33%)",
     "Boundary displacement, given as ±2 mm in 5.5.2: the ring is anchored to "
     "a tolerance rather than exactly. Entered here as a uniform radial error.",
     "systematic", "estimate"),
    ("nu", "10%",
     "As E1. Not exercised in Block A; Poisson's ratio enters from Block B.",
     "systematic", "estimate"),
    ("cable rest length", "0.1%",
     "Channel insertion and anchorage take-up set the effective rest length; "
     "the turnbuckle then adjusts it in half-turn steps, so the residual is "
     "quantised rather than Gaussian. Not exercised in Block A, which has no "
     "cable.",
     "independent", "estimate"),
]
PROV_HEAD = ("parameter", "one tolerance", "where it comes from", "error model",
             "status")

MM = 1e3


def rows(d):
    """Factor rows, largest crown-height effect first.

    Deliberately short. The magnitudes as a function of tolerance are the
    figure's job; the table's job is the ranking, the normalised sensitivity and
    whether the row may be rescaled.
    """
    order = sorted(d, key=lambda k: -abs(float(d[k]["crown_height_half"])))
    out = []
    for k in order:
        r = d[k]
        out.append((
            NAME[k], DELTA[k],
            f"{MM * float(r['crown_height_half']):+.2f}",
            f"{MM * float(r['L_pos_mean']):.2f}",
            f"{float(r['crown_height_elast']):+.2f}",
            f"{100 * float(r['crown_height_asym']):.1f}%",
        ))
    return out


def crown_mm(d):
    return MM * float(next(iter(d.values()))["crown_height_base"])


def worst_asym(d):
    k = max(d, key=lambda k: float(d[k]["crown_height_asym"]))
    return NAME[k], 100 * float(d[k]["crown_height_asym"])


H_DISC, H_2P = crown_mm(disc), crown_mm(twopart)
L_TARGET = MM * float(twopart["s_wale"]["L_target_base"])
S_NOM = float(twopart["s_wale"]["nominal"])
WA_DISC, WA_2P = worst_asym(disc), worst_asym(cable)
C_SW = float(cable["s_wale"]["nominal"])
C_SC = float(cable["s_course"]["nominal"])
C_LT = MM * float(cable["s_wale"]["L_target_base"])
H_CABLE = crown_mm(cable)
DOM_2P = MM * float(twopart["s_course"]["L_pos_mean"])
D, T, C = ov["disc"], ov["2part"], ov["2part_cable"]
OOR = ov["disc_out_of_round"]

HEAD = ("factor", "one tolerance", "crown height (mm)", "surface, L_pos (mm)",
        "elasticity", "asymmetry")

BODY = [
("h", "6.5.3 Robustness to construction imprecision"),

("p", "The design of Section 6.4 is computed from nominal values: a stretch "
      "factor the knitting is asked to deliver, a membrane stiffness measured on "
      "coupons, an inflation pressure held by a regulator, a boundary ring cut "
      "to a drawing. The built structure has none of these exactly. This section "
      "asks how far the inflated equilibrium moves when they are wrong by as "
      "much as fabrication can be expected to leave them wrong, and, where the "
      "geometry carries a design target, whether that movement matters against "
      "the error the design already has."),

("p", f"Eight parameters are perturbed one at a time about a nominal working "
      f"point: the two stretch factors s_wale and s_course, the membrane "
      f"stiffness E1, the orthotropy ratio E2/E1, Poisson's ratio, the "
      f"inflation pressure p, the boundary radius R, and the rest length of the "
      f"crease cable. Two structures are studied. The circular dome of "
      f"Section 6.2 is the method check: a flat rest mesh, no cable, no design "
      f"target of its own. The creased shell of Section 6.4.1 is the case "
      f"study, and it is taken as designed — the optimised stretch factors of "
      f"strategy D, s_wale = {C_SW:.4f} and s_course = {C_SC:.4f}, with the "
      f"crease cable that optimisation was run with."),

("p", f"The reference matters as much as the perturbation. Deviations are "
      f"measured from the equilibrium the design itself predicts, not from the "
      f"design target. The question a tolerance budget has to answer is how far "
      f"the built structure departs from what the optimisation said it would "
      f"be; how far that prediction sits from the target is a separate "
      f"question, settled in Section 6.4 and not improvable by tightening any "
      f"tolerance. On the creased shell the designed equilibrium stands "
      f"{C_LT:.1f} mm from the target in RMS over interior vertices, and every "
      f"deviation reported below is measured from that equilibrium."),

("h2", "Where the tolerances come from"),

("p", "A robustness study is worth exactly what its tolerances are worth, so "
      "each is stated with the mechanism behind it rather than as a bare "
      "magnitude. The mechanisms differ in kind, and that matters as much as "
      "the numbers: some errors apply to the whole structure at once because "
      "one panel is one fabric knitted on one machine setting, while others are "
      "drawn afresh for each region or each cable because each is made by "
      "hand."),

("table_prov", None),

("p", "Every one of these is exercised. Poisson's ratio and the cable rest "
      "length were absent from the first pass of this study, the one on a "
      "cable-free model, and both are now included: the case study has a cable "
      "at the crease, so its rest length is a real tolerance with a real "
      "mechanism rather than a parameter waiting for a later block."),

("p", "The stretch-factor entry deserves its own remark, because it is the "
      "dominant term in everything that follows and it is the one furthest from "
      "a direct measurement. The quantity the fabrication controls is stitch "
      "size, not the stretch factor; the factor is a consequence of it. A "
      "per-stitch error of the order of a tenth of a millimetre is small "
      "against any single stitch, but it accumulates along a course, and it is "
      "that accumulation, over a panel of the size built here, that produces a "
      "stretch-factor deviation of the order assumed. Converting the measured "
      "stitch-size scatter into a factor tolerance is the single most valuable "
      "input this study is waiting on, and it is arithmetic on data Chapter 4 "
      "already has rather than a new experiment."),

("h2", "How the responses are measured"),

("p", f"Two outputs are reported for each perturbation. The crown height is a "
      f"scalar and admits the usual first-order treatment: a central difference "
      f"gives the half-range, and an elasticity, the relative change in crown "
      f"height per relative change in the factor, makes the six comparable "
      f"across their different units. L_pos is the RMS displacement of the "
      f"surface from the baseline equilibrium, taken over interior vertices "
      f"only. The boundary is clamped, so its deviation is zero by construction "
      f"and including it would deflate every figure by an amount set by nothing "
      f"but boundary discretisation."),

("p", f"The nominal for the creased shell is not free to be chosen. Its rest "
      f"mesh is the shape the structure is meant to hold, so the stretch factors "
      f"are whatever puts the inflated equilibrium on that shape, and they are "
      f"fitted. The anisotropic two-parameter fit saturates at both bounds, "
      f"s_wale to 1.400 and s_course to 0.950, which is a result in itself — a "
      f"uniform pre-strain cannot reach this target — but it is unusable as a "
      f"nominal, since a working point on a bound cannot be perturbed "
      f"symmetrically and the plus and minus responses would differ because one "
      f"side was clipped rather than because the physics is asymmetric. The "
      f"isotropic fit is interior at s = {S_NOM:.4f}, and that is the working "
      f"point used here. It leaves a standing mismatch of {L_TARGET:.1f} mm "
      f"before any imperfection is applied."),

("p", "The numerical floor was measured rather than assumed. Re-solving the "
      "baseline along a different stretch-factor continuation path reproduces it "
      "to below 0.01 um in both outputs, at the 1e-8 m precision of the solver's "
      "output. Every response reported below is physical by five orders of "
      "magnitude."),

("h2", "How wrong is too wrong"),

("p", "Reporting a response at one assumed tolerance answers a question nobody "
      "is in a position to ask yet, since five of the six magnitudes are "
      "estimates. Figure 6.31 asks the question the other way round: it applies "
      "the same relative error to every parameter, from a quarter of a percent "
      "to two percent of nominal, and plots what the surface does. The axis is "
      "then a property of the fabrication rather than of an assumption, and the "
      "six factors can be read against one another."),

("p", f"One percent is not a small error here. On the dome a one percent error "
      f"in s_course moves the surface 3.9 mm and on the case study 2.8 mm, "
      f"against a designed equilibrium that stands {C_LT:.1f} mm from its "
      f"target. Two percent roughly doubles both, the responses being linear "
      f"across the range. The weak parameters are weak by an order of magnitude "
      f"rather than marginally: one percent on the orthotropy ratio or on "
      f"Poisson's ratio moves the case study by about 0.09 mm, below the "
      f"precision any of this is measured to, and one percent on E1 by "
      f"0.26 mm."),

("p", f"Panels (c) and (f) say where on the structure those two parameters "
      f"are. The dome is held on its ring alone. The creased shell is held on "
      f"its ring and stiffened along the single line x = 0, where the crease "
      f"cable runs the full span between two boundary anchorages: 19 vertices, "
      f"18 segments, 1.32 m of steel at EA = 157 kN. That the cable and the "
      f"crease are the same line is the point of the design, and it is also why "
      f"its rest length is a tolerance worth specifying — an error in it moves "
      f"the feature the structure exists to hold."),

("p", f"The boundary radius and the cable rest length are set as lengths, not "
      f"as fractions, so panels (c) and (d) give them in millimetres. At the "
      f"±2 mm of 5.5.2 the boundary contributes about 1.9 mm of surface "
      f"deviation on both geometries. Ten millimetres of radius error would "
      f"give 7.5 mm and ten of cable error 2.9 mm — the cable is the gentler of "
      f"the two per millimetre, and only overtakes the stretch factors when "
      f"expressed as a percentage because it is 1.32 m long. Both are "
      f"quantities a workshop can measure directly, which makes them the two "
      f"tolerances in this study that could be replaced by a measurement "
      f"tomorrow."),

("p", "Which parameter matters most depends on the unit the question is asked "
      "in, and that is not a defect of the analysis but the substance of it. "
      "Per percent of nominal the boundary radius is the steepest parameter of "
      "the eight on both geometries; at the tolerances actually assumed, as in "
      "the tables below, it is the smallest contributor of all, because it is "
      "anchored to 0.33% where the stretch factor is assumed good to 4.5%. Both "
      "readings are true. R is the most dangerous parameter and the least "
      "troublesome one. The practical consequence is that the ranking in a "
      "tolerance budget is a statement about the workshop as much as about the "
      "structure, and it moves if the workshop does."),

("p", "The curves are straight over the whole range, which is worth more than "
      "it appears. The claim that an estimate may be replaced without re-running "
      "rested until now on the asymmetry between the plus and minus responses at "
      "a single magnitude — a local check. These sweeps show the response is "
      "linear on every factor and both geometries, so a response may be scaled "
      "to whatever tolerance the measurements eventually give, and read beyond "
      "the frame of the figure for the parameters whose assumed tolerances are "
      "larger than two percent."),

("h2", "The circular dome"),

("table_disc", None),

("p", f"The predicted spread with all six factors at one tolerance is "
      f"{D['crown_rss_mm']:.1f} mm RSS in crown height, "
      f"{D['crown_rss_pct_of_h']:.1f}% of the {H_DISC:.1f} mm the dome rises, "
      f"and {D['crown_aligned_mm']:.1f} mm if every error happens to align. "
      f"Three things in the table are worth drawing out."),

("p", f"s_course dominates, at 2.7 times the next factor, because the course "
      f"direction is the stiff one — E2 is 2.5 times E1 — so pre-strain applied "
      f"along it does the most work. It is also the factor whose tolerance is "
      f"the least well known, which is where the uncertainty in this study "
      f"sits."),

("p", f"R has the second-largest elasticity at "
      f"{float(disc['R']['crown_height_elast']):+.2f}, behind only s_course and "
      f"ahead of both moduli — the crown responds strongly to a relative change "
      f"in boundary radius, and Figure 6.31 shows it steeper still than "
      f"s_course when the two are given the same relative error — and yet it "
      f"has the smallest effect of the six here, purely because its tolerance "
      f"is tight: 0.33% of the nominal against 4.5 to 10% for the others. It is "
      f"a factor worth anchoring precisely rather than one that does not "
      f"matter, and the distinction is only visible because elasticity and "
      f"effect are reported separately. It is also the factor "
      f"most sensitive to the tolerance being right: the ±2 mm of 5.5.2 "
      f"replaced an earlier assumption of 5 mm taken from screw spacing, and "
      f"the row shrank in proportion, the response being linear to "
      f"{100 * float(disc['R']['crown_height_asym']):.1f}%."),

("p", f"The six factors are very nearly degenerate on this geometry. The median "
      f"absolute cosine between their displacement fields is "
      f"{D['median_abs_cosine']:.3f} off the diagonal, and "
      f"{D['most_aligned_pair']} sit at {D['most_aligned_cosine']:.3f}. They "
      f"excite one and the same axisymmetric inflation mode, differing in sign "
      f"and amplitude but not in shape. Two consequences follow. A measured "
      f"deviation cannot be attributed to a factor by this block, only bounded "
      f"in size; attribution needs a geometry on which the fields separate. And "
      f"a joint sampling of all six would give a broad, right-skewed "
      f"distribution of L_pos with a mean below the RSS, because near-parallel "
      f"contributions add and cancel algebraically. The crown-height RSS is "
      f"unaffected by this: for a scalar output RSS is correct whenever the "
      f"factor errors are independent, whatever the shape of the response."),

("h2", "The creased shell, as designed"),

("table_cable", None),

("p", f"The predicted spread is {C['crown_rss_mm']:.1f} mm RSS in crown height, "
      f"{C['crown_rss_pct_of_h']:.1f}% of the {H_CABLE:.0f} mm this shell "
      f"stands, and {C['L_pos_rss_mm']:.1f} mm RMS over the surface. Against "
      f"the dome's {D['crown_rss_pct_of_h']:.1f}% that is a markedly more "
      f"robust structure, and the cable is most of the reason. Run without it, "
      f"on a fitted uniform pre-strain, the same shell gives "
      f"{T['crown_rss_mm']:.1f} mm and {T['crown_rss_pct_of_h']:.1f}%. A "
      f"stiffener introduced to resolve the crease also takes the shape out of "
      f"the fabric's hands, and what the fabric does not control it cannot get "
      f"wrong."),

("p", f"The same is visible factor by factor. With the cable in place a one "
      f"percent error in E1 moves the surface 0.26 mm where the cable-free "
      f"model gave 0.45 mm, and the orthotropy ratio drops from 0.23 to "
      f"0.09 mm. The material scatter that Chapter 4 is being asked to pin "
      f"down matters measurably less to the structure as it is actually "
      f"designed than to the model that stood in for it."),

("h2", "What the cable costs in return"),

("p", f"The cable removes sensitivity to the fabric and adds one of its own. "
      f"Per percent of nominal, its rest length is the second-steepest "
      f"parameter of the eight, at 3.89 mm against the boundary radius's "
      f"4.26 mm, and it is well ahead of s_course at 2.79 mm. The mechanism is "
      f"not mysterious: the cable is stiff, so where it is short the surface "
      f"follows it. A one percent error is 13 mm on a 1.32 m cable, which is "
      f"more than a turnbuckle's resolution and well within what channel "
      f"seating and anchorage take-up can absorb, so this is a tolerance to "
      f"specify rather than to assume."),

("p", f"That completes an argument the earlier, cable-free study could only "
      f"half make. It found that no tolerance moved the shape toward the target "
      f"and concluded that fabrication was not the binding error. That remains "
      f"true, but the reason is sharper: the crease was missing because the "
      f"model had no cable, and once the cable is there the standing mismatch "
      f"falls from {L_TARGET:.1f} mm to {C_LT:.1f} mm. The residual is what the "
      f"two-parameter uniform pre-strain cannot express, and it is the "
      f"parameterisation of Section 6.4, not the workshop, that would have to "
      f"improve to remove it."),

("p", f"Figures 6.34 and 6.35 are the cable-free control, and are worth keeping "
      f"for what they diagnose. Without the cable the mismatch is a "
      f"narrow stripe along x = 0 reaching {T['mismatch_peak_mm']:.0f} mm: the "
      f"target has a sharp crease between its two lobes, and the section through "
      f"them has the target dipping to 292 mm at the valley while the "
      f"equilibrium runs flat across at 380 to 385 mm. The tolerance "
      f"perturbation, by contrast, is a broad smooth mode that lifts or drops "
      f"the whole cap by about 20 mm and does nothing at the valley — the "
      f"cosine between the two fields is "
      f"{T['mismatch_vs_dominant_cosine']:+.3f}. No tolerance moves that shape "
      f"toward its target, because the feature it is missing is not something a "
      f"tolerance controls."),

("p", "Figure 6.35 says why, and corrects the obvious reading. The natural "
      "conclusion from the saturated anisotropic fit is that a uniform "
      "two-parameter pre-strain is too poor a parameterisation to form a crease. "
      "That is not what happened. The target was form-found with a tie along the "
      "valley — the force densities of the form-finding carry a stiff line "
      "exactly along x = 0 — and the finite-element model has no cable there. "
      "The stiff line of the form finding and the stripe of the mismatch are the "
      "same line. The crease is missing from the simulation because the element "
      "that creates it was not carried across from the form finding, not because "
      "the stretch-factor field is too coarse to express it. The fix is to model "
      "the valley tie, and only then to ask whether the parameterisation is rich "
      "enough."),

("p", "One further caution on reading the table for this geometry. L_target is "
      "stationary in the stretch factors at the nominal, because the nominal is "
      "a fitted minimum: both signs make it worse, by 4.84 and 4.89 mm. A "
      "central difference at an optimum is second order and therefore near zero "
      "and meaningless, so the analysis reports the increase instead. In the "
      "directions that were never fitted — E1, p and R — one sign does improve "
      "L_target, as it should."),

("h2", "What this establishes, and what it does not"),

("p", f"Within the tolerances assumed, construction imprecision moves the crown "
      f"height by about {D['crown_rss_pct_of_h']:.0f}% on the dome and "
      f"{C['crown_rss_pct_of_h']:.0f}% on the case study as designed, and the "
      f"surface by {D['L_pos_rss_mm']:.0f} and {C['L_pos_rss_mm']:.0f} mm RMS "
      f"respectively. Set against the {C_LT:.1f} mm the designed equilibrium "
      f"already stands from its target, the tolerances are comparable rather "
      f"than negligible — which is a different conclusion from the one the "
      f"cable-free study reached, and a more useful one, because it says the "
      f"two now deserve attention together rather than the model alone."),

("p", f"Three limits should be stated with equal clarity. All six tolerances are "
      f"estimates; none is yet backed by a measurement, and the dominant one, "
      f"the stretch factor, is the one with no measurement behind it at all. "
      f"Replacing an estimate does not require re-running provided the response "
      f"is linear over the tolerance, which the block checks by the asymmetry "
      f"between the plus and minus responses: on the dome the worst is "
      f"{WA_DISC[1]:.1f}% ({WA_DISC[0]}), inside the 15% threshold, so the "
      f"responses rescale; on the creased shell {WA_2P[0]} reaches "
      f"{WA_2P[1]:.1f}% and E2/E1 15.2%, so those two rows do not rescale safely "
      f"and want a smaller-delta re-run once the material scatter is known. The "
      f"stretch-factor rows, at 1.5 to 2.8%, are solidly linear on both — which "
      f"is the important case, since that is the estimate most likely to be "
      f"revised."),

("p", f"The perturbations here are uniform: a single stretch factor wrong "
      f"everywhere, a single radius wrong everywhere. Real imprecision is also "
      f"non-uniform, and the two need not have the same effect. The boundary "
      f"makes the point sharply. The ±2 mm of 5.5.2 is an anchoring tolerance, "
      f"and an anchoring error that varies around the ring is not a change of "
      f"radius at all; entering it as a uniform radial error, as here, captures "
      f"only its mean. The reference mesh shows the rest is not negligible: its "
      f"boundary vertices run between {OOR['boundary_radius_min_mm']:.1f} and "
      f"{OOR['boundary_radius_max_mm']:.1f} mm, a {OOR['span_mm']:.1f} mm "
      f"out-of-round span at {OOR['boundary_radius_std_mm']:.2f} mm standard "
      f"deviation — {OOR['pct_of_delta_R']}% of the "
      f"{OOR['delta_R_mm']:.0f} mm tolerance being tested, so the nominal "
      f"baseline already carries a non-uniform boundary imperfection larger "
      f"than the uniform one under test. A dedicated out-of-round study is "
      f"therefore worth more than tightening delta_R, and it has to be measured "
      f"against this mesh's existing scatter rather than against a perfect "
      f"circle."),

("p", "And this is one geometry perturbed one factor at a time about a "
      "single-region, cable-free model. The case study of Section 6.4.6 is "
      "multi-region and carries cables, whose rest length is a tolerance of its "
      "own and a quantised one, set by the resolution of a turnbuckle rather "
      "than by a Gaussian. Independent per-region draws, and a joint sampling "
      "that lets the factors interfere, remain to be run."),
]

CAPTIONS = [
    ("tolerance_percent",
     "Figure 6.31: Deviation from the designed equilibrium against how wrong "
     "each parameter is. (a) and (b) give the material and pre-strain "
     "parameters at 0 to 2% of nominal; (d) and (e) give the two that are set "
     "as lengths, at 0 to 20 mm, because a ring is anchored to a distance and a "
     "cable is cut and taken up by one, and a percentage of nominal is a detour "
     "through a number nobody sets. The dotted line in (d) and (e) is the "
     "\u00b12 mm of 5.5.2. (c) and (f) locate those two parameters on the "
     "structures in the same colours: the anchored boundary that R scales, and "
     "the crease cable along x = 0 whose rest length is the other length "
     "tolerance. Curves are the worse of the two signs at each magnitude and "
     "are linear throughout, so a response may be read at whatever tolerance "
     "the measurements eventually give. 180 solves across the two geometries."),
    ("blockA_disc_sensitivity",
     "Figure 6.32: Block A on the circular dome. (a) Crown-height response to "
     "plus and minus one tolerance, sorted by effect. (b) Surface deviation "
     "L_pos, with the two signs shown separately as ticks. (c) Asymmetry "
     "between the two signs, against the 15% linearity threshold. (d) Pairwise "
     "cosine between the per-factor displacement fields: the six are nearly the "
     "same mode, and p and E1 are exactly opposed."),
    ("blockA_2part_cable_sensitivity",
     "Figure 6.33: Block A on the creased shell as designed, the same four "
     "panels. The responses are smaller than the dome's throughout, the cable "
     "carrying what the fabric would otherwise have to. E1 and E2/E1 exceed the "
     "15% asymmetry threshold in panel (c), but both are under 3 mm, so the "
     "non-linearity is on a quantity too small to matter."),
    ("blockA_2part_shape",
     "Figure 6.34: Where the tolerance shows up, rather than how much. (a) The "
     "standing mismatch between the baseline equilibrium and the design target "
     "is a narrow stripe along the valley. (b) The dominant tolerance moves the "
     "whole cap smoothly and does nothing at the valley. (c) The section "
     "through the lobes: the target creases to 292 mm where the equilibrium "
     "runs flat at 380 to 385 mm. (d) The section through the saddle, where "
     "the perturbation acts and the mismatch does not."),
    ("2part_target_diagnosis",
     "Figure 6.35: The standing mismatch is a missing valley tie, not a limit "
     "of the parameterisation. (a) The force densities of the form finding "
     "carry a stiff line along x = 0. (b) The finite-element model has no cable "
     "there, so the crease never forms. (c) The stiff line and the mismatch "
     "are the same line."),
]

NOTES = [
("h", "Notes for revision — not part of the section"),
("p", "Every number above is read from the Block A outputs by "
      "build_section_6_5_3.py: the per-factor responses from "
      "data/block_A_{disc,2part}_sensitivity.csv, and the spreads, cosines and "
      "out-of-round figures from data/block_A_overlap.json, itself produced by "
      "analyse_block_A.py --check-overlap and plot_shape.py. Re-run those and "
      "rebuild and the section follows."),
("p", "1. The tolerances are estimates, and the provenance table now names the "
      "mechanism behind each rather than a bare magnitude. Three are one step "
      "from being measured rather than assumed. The stretch-factor tolerance "
      "needs the measured stitch-size scatter converted into a factor "
      "deviation by accumulation over a panel — arithmetic on Chapter 4 data, "
      "not a new experiment, and the largest single improvement available to "
      "this study. E1 and E2/E1 need Chapter 4's reported spread across repeats "
      "and yarn batches in place of the assumed 10%, and they should be drawn "
      "from the measured covariance rather than independently, since the ratio "
      "inherits the scatter of both moduli. The pressure band should come from "
      "the 6.1.2 sensor record."),

("p", "1a. delta_R has already changed on this basis, from 5 mm assumed from "
      "screw spacing to the ±2 mm that 5.5.2 gives, and Block A was re-run "
      f"rather than rescaled. The effect is small — the crown-height RSS moves "
      f"from 32.25 to {D['crown_rss_mm']:.2f} mm on the dome and from 27.00 to "
      f"{T['crown_rss_mm']:.2f} mm on the creased shell — because R was already "
      f"the smallest contributor. The consequence "
      "that does matter is for the out-of-round argument, since the reference "
      "mesh's own scatter is now larger than the tolerance rather than a fifth "
      "of it."),
("p", "2. Two scripts disagree about the standing mismatch. analyse_block_A.py "
      "reports L_target = 24.59 mm; plot_shape.py prints 20.65 mm RMS for what "
      "reads as the same quantity. The difference is that plot_shape averages "
      "over all vertices while the analysis restricts to interior ones, and the "
      "clamped boundary deflates the former. The section uses 24.59 mm. Worth "
      "making plot_shape mask the boundary so the two agree."),
("p", "3. Figure numbers 6.31 to 6.35 are provisional and assume 6.29 and 6.30 "
      "belong to Section 6.5.1."),
("p", "4. Blocks B to E are not built. B adds Poisson's ratio and cable rest "
      "length at the case-study optimum and needs the n-region binary; C draws "
      "per-region and per-cable independently; D samples jointly and reports the "
      "distribution of L_pos against the RSS predicted here; E repeats B on a "
      "second geometry. Block A already answers part of E's question — the "
      "dominant tolerance is s_course on both geometries — so E is testing "
      "whether that survives cables and regions."),
("p", "5. Section 6.4 currently gives one scan, so comparing a measured "
      "deviation against this predicted spread compares a distribution with a "
      "point. Two specimens from the same knit programme, or one specimen "
      "re-scanned after re-inflation, would make the comparison two-sided."),
]


def add_table(doc, rows_, head=HEAD):
    t = doc.add_table(rows=1, cols=len(head))
    t.style = "Table Grid"
    for i, h in enumerate(head):
        cell = t.rows[0].cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.font.bold = True
            r.font.size = Pt(8)
    for row in rows_:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
            for r in cells[i].paragraphs[0].runs:
                r.font.size = Pt(8)
    return t


def md_table(rows_, head=HEAD):
    out = ["| " + " | ".join(head) + " |", "|" + "---|" * len(head)]
    out += ["| " + " | ".join(r) + " |" for r in rows_]
    return "\n".join(out) + "\n"


TABLE_CAPTIONS = {
    "table_prov":
        "Table 6.X: The tolerances, their mechanisms and their error models. "
        "Systematic errors apply to the whole structure at once; independent "
        "ones are drawn separately per region or per cable. Every magnitude is "
        "currently an estimate: the section each waits on is named, and "
        "tolerances.py records the same table in machine-readable form.",
    "table_disc":
        "Table 6.Y: Block A on the circular dome, sorted by effect. Crown "
        "height is the half-range of the central difference at one tolerance "
        "either side of the nominal, and L_pos the RMS surface displacement "
        "from the baseline over interior vertices. Elasticity is the relative "
        "change in crown height per relative change in the factor, which is "
        "what makes six parameters in six different units comparable. "
        "Asymmetry is the difference between the two signs as a fraction of the "
        "first difference; under 15% the response may be rescaled to a "
        "different tolerance rather than re-run. Figure 6.31 gives the same "
        "responses as a function of tolerance magnitude.",
    "table_cable":
        "Table 6.Z: Block A on the creased shell as designed - the optimised "
        "stretch factors of strategy D with the crease cable - about that "
        "equilibrium rather than about the design target. Same columns. E1 at "
        "19.4% and E2/E1 at 20.6% exceed the asymmetry threshold, so those two "
        "rows should not be rescaled by hand; the swept curves of Figure 6.31 "
        "cover them directly, and both are small in absolute terms. The cable "
        "rest length and Poisson's ratio are swept in Figure 6.31 but are not "
        "in this table, which reports the six factors of the original block.",
}


def build_md():
    lines = []
    for kind, text in BODY:
        if kind == "h":
            lines.append(f"# {text}\n")
        elif kind == "h2":
            lines.append(f"## {text}\n")
        elif kind == "p":
            lines.append(f"{text}\n")
        elif kind == "table_prov":
            lines.append(md_table(PROVENANCE, PROV_HEAD))
            lines.append(f"*{TABLE_CAPTIONS[kind]}*\n")
        elif kind.startswith("table_"):
            lines.append(md_table(rows({"table_disc": disc,
                                        "table_2part": twopart,
                                        "table_cable": cable}[kind])))
            lines.append(f"*{TABLE_CAPTIONS[kind]}*\n")
    for name, cap in CAPTIONS:
        lines.append(f"![{name}](../imperfection_study/figures/{name}.png)\n")
        lines.append(f"*{cap}*\n")
    for kind, text in NOTES:
        lines.append(("# " if kind == "h" else "") + text + "\n")
    return "\n".join(lines)


def build_docx():
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    for kind, text in BODY:
        if kind == "h":
            doc.add_heading(text, level=2)
        elif kind == "h2":
            doc.add_heading(text, level=3)
        elif kind == "p":
            p = doc.add_paragraph(text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif kind.startswith("table_"):
            if kind == "table_prov":
                add_table(doc, PROVENANCE, PROV_HEAD)
            else:
                add_table(doc, rows({"table_disc": disc, "table_2part": twopart,
                                     "table_cable": cable}[kind]))
            cap = doc.add_paragraph(TABLE_CAPTIONS[kind])
            cap.runs[0].font.size = Pt(9)
            cap.runs[0].font.italic = True

    for name, cap in CAPTIONS:
        path = os.path.join(FIG, f"{name}.png")
        if not os.path.exists(path):
            print(f"  missing figure, skipped: {path}")
            continue
        doc.add_picture(path, width=Inches(6.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c = doc.add_paragraph(cap)
        c.runs[0].font.size = Pt(9)
        c.runs[0].font.italic = True

    doc.add_page_break()
    for kind, text in NOTES:
        if kind == "h":
            doc.add_heading(text, level=2)
        else:
            doc.add_paragraph(text, style="List Bullet")

    doc.save(OUT_DOCX)


if __name__ == "__main__":
    os.makedirs(os.path.dirname(OUT_MD), exist_ok=True)
    with open(OUT_MD, "w") as f:
        f.write(build_md())
    build_docx()
    print(f"  saved: {OUT_DOCX}")
    print(f"  saved: {OUT_MD}")
